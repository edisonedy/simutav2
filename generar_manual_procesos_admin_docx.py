from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("MANUAL_PROCESOS_ADMINISTRATIVOS.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=11, bold=False, color="000000", after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, name="Calibri", size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=11)
    return p


def add_heading(doc, text, level):
    sizes = {1: 18, 2: 14, 3: 12}
    colors = {1: "2E5D7B", 2: "2E5D7B", 3: "1F1F1F"}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])
    return p


def add_info_box(doc, title, lines):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F7FA")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color="1F1F1F")
    for line in lines:
        pp = cell.add_paragraph()
        pp.paragraph_format.space_after = Pt(2)
        rr = pp.add_run(line)
        set_run_font(rr, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.4)
    hdr = table.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Contenido"
    for c in hdr:
        set_cell_shading(c, "E8EEF5")
        for p in c.paragraphs:
            for r in p.runs:
                set_run_font(r, size=10, bold=True)
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for p in cells[0].paragraphs + cells[1].paragraphs:
            for r in p.runs:
                set_run_font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(6)
r = title.add_run("Manual de configuracion\nProcesos administrativos")
set_run_font(r, name="Calibri", size=24, bold=True, color="1F1F1F")

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(18)
r = sub.add_run("Caso ejemplo: Mejora de atencion estudiantil en secretaria academica")
set_run_font(r, size=12, color="555555")

add_info_box(
    doc,
    "Objetivo del documento",
    [
        "Explicar paso a paso como crear y configurar la simulacion.",
        "Dejar un ejemplo completo que pueda replicarse en otro servidor.",
        "Usar solo la configuracion necesaria para trabajar con IA.",
    ],
)

add_heading(doc, "1. Flujo general", 1)
for item in [
    "Crear una simulacion nueva.",
    "Llenar lo basico del caso.",
    "Entrar a Configurar.",
    "Completar Caso.",
    "Completar Evaluacion con IA.",
    "Agregar opciones avanzadas solo si hacen falta.",
    "Publicar la simulacion.",
]:
    add_number(doc, item)

add_heading(doc, "2. Datos base del caso", 1)
add_two_col_table(
    doc,
    [
        ("Titulo", "Mejora de atencion estudiantil en secretaria academica"),
        ("Tema", "Procesos administrativos"),
        ("Rol estudiante", "Asistente administrativo de coordinacion academica"),
        (
            "Contexto",
            "La secretaria academica recibe solicitudes de matriculas, certificados, cambios de paralelo y correcciones de notas. La atencion es lenta, desordenada y genera acumulacion de tramites.",
        ),
        (
            "Objetivo",
            "Organizar la atencion de los tramites para reducir pendientes, mejorar el orden del proceso y aumentar la satisfaccion estudiantil.",
        ),
        (
            "Situacion inicial",
            "En la primera semana del periodo academico se acumulan solicitudes y no existe un orden claro de atencion. El estudiante debe proponer como organizar el trabajo.",
        ),
        ("Modo de simulacion", "Con IA - Simulacion dinamica"),
        ("Rondas", "3"),
        ("Tiempo estimado", "30 minutos"),
        ("Nivel de dificultad", "Media"),
    ],
)

add_heading(doc, "3. Como llenar Nueva simulacion", 1)
for item in [
    "Entrar a Simulador.",
    "Hacer clic en Nueva simulacion.",
    "Elegir la materia Procesos Administrativos.",
    "Elegir el modo Con IA - Simulacion dinamica.",
    "Escribir el titulo, tema, rol, contexto, objetivo y situacion inicial.",
    "Guardar para pasar a Configurar.",
]:
    add_number(doc, item)

add_info_box(
    doc,
    "Importante",
    [
        "En esta pantalla no hace falta llenar configuraciones tecnicas.",
        "Lo importante es dejar bien escrito el caso.",
        "La IA se apoya sobre todo en la rubrica y los indicadores.",
    ],
)

add_heading(doc, "4. Configurar: solo lo necesario", 1)
add_para(doc, "La pantalla Configurar se divide en tres bloques:", bold=True, after=4)
for item in [
    "Caso",
    "Evaluacion con IA",
    "Opciones avanzadas",
]:
    add_bullet(doc, item)

add_heading(doc, "5. Caso", 2)
add_para(doc, "En este bloque se define lo que el estudiante va a leer.", after=4)
for item in [
    "Caso y aprendizaje: contexto, objetivo y situacion inicial.",
    "Datos visibles del caso: opcional. Sirve para mostrar tablas, alternativas o informacion de apoyo.",
]:
    add_bullet(doc, item)

add_heading(doc, "6. Evaluacion con IA", 2)
add_para(doc, "Esta parte define como la IA revisa la respuesta del estudiante.", after=4)
for item in [
    "Indicadores: muestran si la situacion mejora o empeora.",
    "Conceptos esperados por ronda: son la rubrica que define la nota.",
]:
    add_bullet(doc, item)

add_info_box(
    doc,
    "Como leer los indicadores",
    [
        "+5 significa que el indicador sube.",
        "-5 significa que el indicador baja.",
        "Eso puede ser bueno o malo dependiendo de si el indicador conviene alto o bajo.",
    ],
)

add_heading(doc, "7. Indicadores del caso", 1)
for item in [
    "Orden del proceso",
    "Tramites pendientes",
    "Solicitudes atendidas a tiempo",
    "Satisfaccion estudiantil",
    "Errores o reprocesos",
]:
    add_bullet(doc, item)

add_heading(doc, "8. Rubrica completa por rondas", 1)

add_heading(doc, "Ronda 1 - Diagnostico del problema", 2)
for item in [
    "Identifica desorden en la atencion - Peso 30 - Critico",
    "Reconoce falta de prioridades - Peso 25 - Critico",
    "Detecta falta de seguimiento - Peso 25 - Critico",
    "Menciona impacto en estudiantes - Peso 20 - Critico",
]:
    add_bullet(doc, item)

add_heading(doc, "Ronda 2 - Decision / propuesta de organizacion", 2)
for item in [
    "Clasifica los tipos de tramite - Peso 25 - Critico",
    "Define prioridades de atencion - Peso 30 - Critico",
    "Asigna responsables - Peso 25 - Critico",
    "Propone un registro de seguimiento - Peso 20 - No critico",
]:
    add_bullet(doc, item)

add_heading(doc, "Ronda 3 - Plan / control y mejora", 2)
for item in [
    "Define tiempos maximos de respuesta - Peso 25 - Critico",
    "Propone indicadores de control - Peso 30 - Critico",
    "Evalua reduccion de pendientes - Peso 20 - Critico",
    "Justifica la mejora del proceso - Peso 25 - No critico",
]:
    add_bullet(doc, item)

add_info_box(
    doc,
    "Impacto sugerido para los conceptos clave",
    [
        "Si cumple: Orden del proceso +5 y Tramites pendientes -5.",
        "Si falta: Orden del proceso -5 y Tramites pendientes +5.",
        "Los demas indicadores pueden quedar en 0 si no se quiere moverlos desde ese concepto.",
    ],
)

add_heading(doc, "9. Opciones avanzadas", 1)
for item in [
    "Restricciones: opcionales. Sirven para penalizar si un indicador queda en mala zona.",
    "Recursos: opcionales. Sirven para representar tiempo, presupuesto o capacidad limitada.",
    "Decisiones sugeridas: opcionales. Son ejemplos que el estudiante puede elegir.",
    "Eventos dinamicos: opcionales. Son cambios que pueden aparecer durante las rondas.",
]:
    add_bullet(doc, item)

add_heading(doc, "10. Cuando la simulacion ya se puede publicar", 1)
for item in [
    "Tiene titulo.",
    "Tiene contexto.",
    "Tiene objetivo.",
    "Tiene situacion inicial.",
    "Tiene indicadores.",
    "Tiene conceptos esperados por ronda.",
    "Cada ronda suma 100 en la rubrica.",
]:
    add_bullet(doc, item)

add_heading(doc, "11. Como leer el resultado del estudiante", 1)
for item in [
    "Puntaje academico: mide que tan bien cumplio la rubrica.",
    "Salud del caso: mide como quedaron los indicadores al final.",
    "Un estudiante puede sacar 100 en puntaje academico y aun asi dejar indicadores en estado medio o bajo.",
]:
    add_bullet(doc, item)

add_heading(doc, "12. Pasos para replicar en otro servidor", 1)
for item in [
    "Crear la simulacion con los datos base del caso.",
    "Configurar los 5 indicadores listados en este manual.",
    "Cargar los conceptos de la ronda 1, 2 y 3 con sus pesos.",
    "Revisar que cada ronda sume 100.",
    "Probar una ejecucion como estudiante.",
    "Publicar cuando la rubrica e indicadores ya esten completos.",
]:
    add_number(doc, item)

add_heading(doc, "13. Resumen corto", 1)
add_para(
    doc,
    "Para replicar este caso en otro servidor no hace falta configurar todo desde el inicio. Lo esencial es crear bien el caso, definir indicadores claros y cargar la rubrica completa por rondas.",
    after=6,
)

doc.save(OUT)
print(OUT.resolve())
